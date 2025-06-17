import streamlit as st
import inference as inf
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import os
from transformers import AutoModelForCausalLM,AutoTokenizer,BitsAndBytesConfig,TextIteratorStreamer
import io
import tempfile
import subprocess
import threading
def chat1(model,tok,question=None):
    if(question==None):
        question=input("Give your question :: ")
    message=[
        {
            "role": "system",
"content": f"""
You are a helpful assistant that provides concise, direct responses and can generate high-quality, detailed documents when asked.

**Classification Steps:**

**Step 1: MULTIQUERY Check**
    - If the user asks to perform more than one task in a single prompt (e.g., "Make a PDF on X and a Word on Y" or "Generate a DOCX and also explain Y"), classify as MULTIQUERY.
    - If MULTIQUERY is detected, STOP here and return:
    - A Python list of strings (STRICTLY),
    - Each string should clearly describe one task or document, such as:
        "Generate a DOCX file having information about the solar system."
    - Wrap the list inside a single <multiquery> tag
    - Do NOT proceed to Step 2 or Step 3

**Step 2: DOCUMENT_GENERATION Check**
    1. Create a relevant title wrapped in <title></title>
    2. Generate **rich, informative, multi-paragraph content** that:
        - Covers the topic in depth,
        - Flows logically with clear sections (e.g., Introduction, Details, Examples, Conclusion),
        - Is **at least one page worth of text**,
        - Should be Professionally Written
        - DON'T OVERTHINK AND RESPOND AS FAST AS POSSIBLE
    3. End with exactly one tag:
        - <DownloadPDF> for PDF requests
        - <DownloadDOCX> for Word requests
    4. Include ONLY: properly structured title, content, and tag

**Step 3: REGULAR_QUERY**
    - If it is not a document generation task, respond normally.
    - Do NOT return <title>, download tags, or <multiquery>

**Important Rules:**
- Follow classification steps in order.
- If MULTIQUERY is true, do NOT go to Step 2 or Step 3.
- For MULTIQUERY, never return title, content, or tag fields — only the list of user-style instructions.
- When generating a document, aim for **depth, clarity, and structure** to produce a meaningful document, not just a short summary.

"""
        },
        {
            "role":"user",
            "content": question
        }
    ]

    prompt=tok.apply_chat_template(message,tokenize=False)
    # print(prompt)
    tok_prompt=tok(prompt,return_tensors='pt').to('cuda')
    
    streamer = TextIteratorStreamer(tok, skip_prompt=True, skip_special_tokens=True)
    generation_kwargs = dict(tok_prompt, streamer=streamer, max_new_tokens=4096)
    thread = threading.Thread(target=model.generate, kwargs=generation_kwargs)
    thread.start()
    output_container = st.empty()
    full_output = ""
    for token in streamer:
        # print(token, end="", flush=True)
        full_output += token
        output_container.markdown(full_output)
    return full_output
    # return "Streaming done"
@st.cache_data
def doc_creation(out, title):
    docx_file = f"{title}.docx"
    doc = Document()
    doc.add_heading(title, 0)

    for line in out.strip().split("\n")[0:]:  # Skip title and blank
        if line.startswith("*"):
            para = doc.add_paragraph(line.replace("*", u"\u2022"), style='List Bullet')
        elif line.strip() == "<Download>":
            doc.add_paragraph("<Download>").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            doc.add_paragraph(line.strip())
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer.getvalue()

    
def create_pdf_bytes(out, title):
    if out is None:
        return None
    
    # Create temporary directory for file operations
    with tempfile.TemporaryDirectory() as temp_dir:
        docx_file = os.path.join(temp_dir, f"{title}.docx")
        pdf_file = os.path.join(temp_dir, f"{title}.pdf")
        
        # Create Word document (same as your original logic)
        doc = Document()
        doc.add_heading(title, 0)

        for line in out.strip().split("\n")[0:]:
            if line.startswith("*"):
                para = doc.add_paragraph(line.replace("*", u"\u2022"), style='List Bullet')
            elif line.strip() == "<Download>":
                doc.add_paragraph("<Download>").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                doc.add_paragraph(line.strip())

        # Save Word document
        doc.save(docx_file)
        subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", temp_dir,
            docx_file
        ], check=True, timeout=60)

        # Read PDF file and convert to bytes
        with open(pdf_file, 'rb') as pdf_f:
            pdf_bytes = pdf_f.read()
        
        return pdf_bytes
            


def clean_output(text):
    flag=-1
    ind1=text.find("<think>")
    ind2=text.find("</think>")
    title=None
    text_out=None
    # print(text[ind1+len("<think>"):ind2])
    if re.search(r"<DownloadPDF>(?![>\w])", text):
        flag=1
    elif re.search(r"<DownloadDOCX>(?![>\w])", text):
        flag=0
    elif re.search(r"</multiquery>(?![>\w])", text):
        flag=2
    if(flag!=-1):
        ind3=text[ind2:].find("<title>")
        ind4=text[ind2:].find("</title>")
        title=text[ind2+ind3+len("<title>"):ind2+ind4]
        text_out=text[ind2+ind4+len('</title>'):]
    else:
        text_out=text[ind2+len("</think>"):]
    think=text[ind1:ind2]
    return flag,think,title,text_out

def give_model():
    bnb_config = BitsAndBytesConfig(
    load_in_4bit=True,
    bnb_4bit_use_double_quant=True,
    bnb_4bit_quant_type="nf4",
    bnb_4bit_compute_dtype="float16"
    )
    tok=AutoTokenizer.from_pretrained(
        "deepseek-ai/DeepSeek-R1-0528-Qwen3-8B"
    )
    model=AutoModelForCausalLM.from_pretrained(
        "deepseek-ai/DeepSeek-R1-0528-Qwen3-8B",
        device_map='auto',
        quantization_config=bnb_config,
        # torch_dtype=torch.float32
    )
    return model,tok

st.title("Chatbot-Interface")
if "Response" not in st.session_state:
    st.session_state.Response=[]
    st.session_state.model_response=[]
# print(st.session_state)
for response in range(len(st.session_state.Response)):
    st.markdown(st.session_state.Response[response])
    st.markdown(st.session_state.model_response[response])
# chat=st.chat_input("Message...")
if 'files' not in st.session_state:
    st.session_state.files=[]
if "de" not in st.session_state:
    st.session_state.de=[]
if(len(st.session_state.de)!=0):
    for i in st.session_state.de:
        try:
            st.session_state.de.pop()
            # st.session_state.files.pop()
        except:
            pass

if(len(st.session_state.files)!=0):
    with st.sidebar:
        st.title("Your Files")
        for i in range(len(st.session_state.files)):
            if(st.session_state.files[i][1]=='pdf'):
                st.session_state.de.append(st.download_button(
                    f"{st.session_state.files[i][0]} download button",
                    data=create_pdf_bytes(st.session_state.files[i][2],st.session_state.files[i][0]),
                    file_name=f"{st.session_state.files[i][0]}.pdf",
                    key=i
                ))
            else:
                st.session_state.de.append(st.download_button(
                    f"{st.session_state.files[i][0]} download button",
                    data=doc_creation(st.session_state.files[i][2],st.session_state.files[i][0]),
                    file_name=f"{st.session_state.files[i][0]}.docx",
                    key=i
                ))

if 'model' not in st.session_state:
    st.session_state.model = None
if 'tok' not in st.session_state:
    st.session_state.tok = None

# Load model only once
if st.session_state.model is None:
    with st.spinner("Loading model... This may take a few minutes."):
        st.session_state.model, st.session_state.tok = give_model()
    st.success("Model loaded successfully!")

# Use the model
model = st.session_state.model
tok = st.session_state.tok
chat=st.chat_input("Message...")
# def trunc_history()

# def generate_history():
#     history=[]
#     for i in range(len(st.session_state.Response)):
#         di1={"role":"user","content":st.session_state.Response[i]}
#         di2={"role":'assistant',"content":st.session_state.model_response[i]}
#         history.append(di1)
#         history.append(di2)
#     return history


def return_list(text):
    ind1=text.find("<multiquery>")
    ind2=text.find("</multiquery")
    return eval(text[ind1+len("<multiquery>"):ind2].strip())

def output(model,tok,question=None):
    # hist1=generate_history()
    out=chat1(model,tok,question)
    flag,think,title,text=clean_output(out)
    try:
        if(flag==-1):
            print("Flag",flag)
            st.markdown(text)
        elif(flag==2):
            li=return_list(text)
            for i in li:
                output(model,tok,i)
                # st.session_state.files.pop()
        else:
            st.markdown("Word" if flag==0 else "PDF")
            st.markdown("-"*50)
            st.markdown(think)
            st.markdown("-"*50)
            st.markdown(title)
            st.markdown("-"*50)
            st.markdown(text)
            st.markdown('-'*50)
            if(flag==0):
                st.session_state.files.append([title,'docx',text])
        #     if(flag==0):
        #         st.session_state.files.append(st.download_button(
        #             label=f"{title} download button",
        #             data=doc_creation(text,title),
        #             file_name=f"{title}.docx",
        #             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        #         ))
            elif(flag==1):
                st.session_state.files.append([title,'pdf',text])
        #         st.session_state.files.append(st.download_button(
        #             label=f"{title} download button",
        #             data=create_pdf_bytes(text,title) ,
        #             file_name=f"{title}.pdf",
        #             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        #         ))
        st.session_state.Response.append(question)
        st.session_state.model_response.append(text)
    except Exception as e:
        print("Error")
        print(e)
        st.markdown("Word" if flag==0 else "PDF")
        st.markdown("-"*50)
        st.markdown(think)
        st.markdown("-"*50)
        st.markdown(title)
        st.markdown("-"*50)
        st.markdown(text)
        st.markdown('-'*50)
if 'chats' not in st.session_state:
    st.session_state.chats=""
# print("*"*200)
# print(chat)
# print(st.session_state.chats)
# print("*"*200)
if(chat!=None or st.session_state.chats!=""):
    if(chat !=None):
        st.markdown(chat)
        st.session_state.chats=chat
        st.rerun()
    else:
        st.markdown(st.session_state.chats)
    
    
    st.session_state.files=[]
    # hist=generate_history()
    output(model,tok,st.session_state.chats)
    print(len(st.session_state.files))
    if(len(st.session_state.files)!=1):
        with st.sidebar:
            st.title("Your Files")
            for i in range(len(st.session_state.files)):
                if(st.session_state.files[i][1]=='pdf'):
                    st.session_state.de.append(st.download_button(
                        f"{st.session_state.files[i][0]} download button",
                        data=create_pdf_bytes(st.session_state.files[i][2],st.session_state.files[i][0]),
                        file_name=f"{st.session_state.files[i][0]}.pdf",
                        key=i
                    ))
                else:
                    st.session_state.de.append(st.download_button(
                        f"{st.session_state.files[i][0]} download button",
                        data=doc_creation(st.session_state.files[i][2],st.session_state.files[i][0]),
                        file_name=f"{st.session_state.files[i][0]}.docx",
                        key=i
                    ))
    st.session_state.chats=""
    st.rerun()
