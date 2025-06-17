from transformers import AutoModelForCausalLM,AutoTokenizer,BitsAndBytesConfig,TextStreamer
import torch
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert
import os
import re
import subprocess

# import st

def create_pdf_bytes(out, title):
    if out is None:
        return None
    doc = Document()
    doc.add_heading(title, 0)

    for line in out.strip().split("\n")[2:]:
        if line.startswith("*"):
            para = doc.add_paragraph(line.replace("*", u"\u2022"), style='List Bullet')
        elif line.strip() == "<Download>":
            doc.add_paragraph("<Download>").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            doc.add_paragraph(line.strip())

    # Use temporary files for conversion
    with tempfile.TemporaryDirectory() as temp_dir:
        docx_path = os.path.join(temp_dir, f"{title}.docx")
        pdf_path = os.path.join(temp_dir, f"{title}.pdf")
        
        # Save docx temporarily
        doc.save(docx_path)
        
        # Convert to PDF using LibreOffice
        try:
            subprocess.run([
                "libreoffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", temp_dir,
                docx_path
            ], check=True)
            
            # Read PDF bytes
            with open(pdf_path, 'rb') as pdf_file:
                pdf_bytes = pdf_file.read()
            
            return pdf_bytes
            
        except subprocess.CalledProcessError as e:
            st.error(f"Error converting to PDF: {e}")
            return None
    
# pdf_creation(out)
def doc_creation(out,title):
    docx_file = f"{title}.docx"
    doc = Document()
    doc.add_heading(title, 0)

    for line in out.strip().split("\n")[2:]:  # Skip title and blank
        if line.startswith("*"):
            para = doc.add_paragraph(line.replace("*", u"\u2022"), style='List Bullet')
        elif line.strip() == "<Download>":
            doc.add_paragraph("<Download>").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            doc.add_paragraph(line.strip())
    doc.save(docx_file)
    print("Doc created successfully:", os.path.abspath(docx_file))
def clean_output(text):
    flag=-1
    ind1=text.find("<think>")
    ind2=text.find("</think>")
    title=None
    text_out=None
    # print(text[ind1+len("<think>"):ind2])
    if re.search(r"<DownloadPDF(?![>\w])", text):
        flag=1
    elif re.search(r"<DownloadDOCX(?![>\w])", text):
        flag=0
    
    if(flag!=-1):
        ind3=text[ind2:].find("<title>")
        ind4=text[ind2:].find("</title>")
        title=text[ind2+ind3+len("<title>"):ind2+ind4]
        text_out=text[ind2+ind4+len('</title>'):]
    else:
        text_out=text[ind2+len("</think>"):]
    think=text[ind1:ind2]
    return flag,think,title,text_out
def chat(model,tok,question=None):
    if(question==None):
        question=input("Give your question :: ")
    message=[
        {
            "role": "system",
"content": f"""
You are a helpful assistant that provides concise, direct responses and can generate high-quality, detailed documents when asked.

**Classification Steps:**

**Step 1: MULTIQUERY Check**
    - If the user asks to perform more than one task in a single prompt (e.g., "Make a PDF on X and a Word on Y" or "Generate a DOCX and also explain Y"), classify as MULTIQUERY.
    - If MULTIQUERY is detected, STOP here and return:
    - A Python list of strings (STRICTLY),
    - Each string should clearly describe one task or document, such as:
        "Generate a DOCX file having information about the solar system."
    - Wrap the list inside a single <multiquery> tag
    - Do NOT proceed to Step 2 or Step 3

**Step 2: DOCUMENT_GENERATION Check**
    1. Create a relevant title wrapped in <title></title>
    2. Generate **rich, informative, multi-paragraph content** that:
        - Covers the topic in depth,
        - Flows logically with clear sections (e.g., Introduction, Details, Examples, Conclusion),
        - Is **at least one page worth of text**,
        - Includes headings and formatting hints if needed (e.g., bold, bullet points).
        - Should be Professionally Written
    3. End with exactly one tag:
        - <DownloadPDF> for PDF requests
        - <DownloadDOCX> for Word requests
    4. Include ONLY: properly structured title, content, and tag

**Step 3: REGULAR_QUERY**
    - If it is not a document generation task, respond normally.
    - Do NOT return <title>, download tags, or <multiquery>

**Important Rules:**
- Follow classification steps in order.
- If MULTIQUERY is true, do NOT go to Step 2 or Step 3.
- For MULTIQUERY, never return title, content, or tag fields — only the list of user-style instructions.
- When generating a document, aim for **depth, clarity, and structure** to produce a meaningful document, not just a short summary.

"""
        },
        {
            "role":"user",
            "content": question
        }
    ]

    prompt=tok.apply_chat_template(message,tokenize=False)
    # print(prompt)
    tok_prompt=tok(prompt,return_tensors='pt').to('cuda')
    
    streamer = TextStreamer(tok, skip_prompt=True, skip_special_tokens=True)
    output=model.generate(**tok_prompt,streamer=streamer,max_new_tokens=4096 )
    out=tok.decode(output[0].tolist(),skip_special_tokens=True)
    idx=out.find("assistant<|end_header_id|>")
    out=out[idx+len("assistant<|end_header_id|>"):].strip()
    out=out[:out.find("<|eot_id|>")]
    return out
    print()
    flag,think,title,text=clean_output(out)
    if(flag==-1):
        print(text)
    else:
        if(flag==0):
            doc_creation(text,title)
        else:
            pdf_creation(text,title)
        print("Word" if flag==0 else "PDF")
        print("-"*50)
        print(think)
        print("-"*50)
        print(title)
        print("-"*50)
        print(text)
        print('-'*50)

if (__name__=='__main__'):
    bnb_config = BitsAndBytesConfig(
    load_in_4bit=True,
    bnb_4bit_use_double_quant=True,
    bnb_4bit_quant_type="nf4",
    bnb_4bit_compute_dtype="float16"
    )
    tok=AutoTokenizer.from_pretrained(
        "deepseek-ai/DeepSeek-R1-0528-Qwen3-8B"
    )
    model=AutoModelForCausalLM.from_pretrained(
        "deepseek-ai/DeepSeek-R1-0528-Qwen3-8B",
        device_map='auto',
        quantization_config=bnb_config,
        # torch_dtype=torch.float32
    )
    while(True):
        chat(model)