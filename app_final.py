import streamlit as st
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import io
import tempfile
import subprocess
import utils as ut

import sys
import logging

# Create logs directory if it doesn't exist



@st.cache_data
def doc_creation(out, title):
    docx_file = f"{title}.docx"
    doc = Document()
    doc.add_heading(title, 0)

    for line in out.strip().split("\n")[0:]:
        if line.startswith("*"):
            para = doc.add_paragraph(line.replace("*", u"\u2022"), style='List Bullet')
        elif line.strip() == "<Download>":
            doc.add_paragraph("<Download>").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            doc.add_paragraph(line.strip())
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer.getvalue()

    
def create_pdf_bytes(out, title):
    if out is None:
        return None
    
    with tempfile.TemporaryDirectory() as temp_dir:
        docx_file = os.path.join(temp_dir, f"{title}.docx")
        pdf_file = os.path.join(temp_dir, f"{title}.pdf")

        doc = Document()
        doc.add_heading(title, 0)

        for line in out.strip().split("\n")[0:]:
            if line.startswith("*"):
                para = doc.add_paragraph(line.replace("*", u"\u2022"), style='List Bullet')
            elif line.strip() == "<Download>":
                doc.add_paragraph("<Download>").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                doc.add_paragraph(line.strip())
        doc.save(docx_file)
        subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", temp_dir,
            docx_file
        ], check=True, timeout=60)
        with open(pdf_file, 'rb') as pdf_f:
            pdf_bytes = pdf_f.read()
        
        return pdf_bytes
try:
    st.title("Chatbot-Interface")
    if 'choice' not in st.session_state:
        st.session_state.choice='Select'
    # st.session_state.choice='Select'
    if st.session_state.choice=='Select':
            st.session_state.choice=st.selectbox("Choose your model ",['Select','DeepSeek','Llama'])
            st.rerun()
    if "Response" not in st.session_state:
        st.session_state.Response=[]
        st.session_state.model_response=[]
    for response in range(len(st.session_state.Response)):
        st.markdown(st.session_state.Response[response])
        st.markdown(st.session_state.model_response[response])
    if 'files' not in st.session_state:
        st.session_state.files=[]
    if "de" not in st.session_state:
        st.session_state.de=[]
    if(len(st.session_state.de)!=0):
        for i in st.session_state.de:
            try:
                st.session_state.de.pop()
            except:
                pass

    if(len(st.session_state.files)!=0):
        with st.sidebar:
            st.title("Your Files")
            for i in range(len(st.session_state.files)):
                if(st.session_state.files[i][1]=='pdf'):
                    st.session_state.de.append(st.download_button(
                        f"{st.session_state.files[i][0]} download button",
                        data=create_pdf_bytes(st.session_state.files[i][2],st.session_state.files[i][0]),
                        file_name=f"{st.session_state.files[i][0]}.pdf",
                        key=i
                    ))
                else:
                    st.session_state.de.append(st.download_button(
                        f"{st.session_state.files[i][0]} download button",
                        data=doc_creation(st.session_state.files[i][2],st.session_state.files[i][0]),
                        file_name=f"{st.session_state.files[i][0]}.docx",
                        key=i
                    ))

    if 'model' not in st.session_state:
        st.session_state.model = None
    if 'tok' not in st.session_state:
        st.session_state.tok = None
    if(st.session_state.choice.lower()!='select'):
        if st.session_state.model is None:
            if(st.session_state.choice=='DeepSeek'):
                with st.spinner("Loading model... This may take a few minutes."):
                    st.session_state.model, st.session_state.tok = ut.give_model_d()
                st.success("Model loaded successfully!")
            elif(st.session_state.choice=='Llama'):
                with st.spinner("Loading model... This may take a few minutes."):
                    st.session_state.model, st.session_state.tok = ut.give_model()
                st.success("Model loaded successfully!")
            else:
                st.info("Please select a model")
        
        
        model = st.session_state.model
        tok = st.session_state.tok
        chat=st.chat_input("Message...")
        if 'chats' not in st.session_state:
            st.session_state.chats=""
        if(chat!=None or st.session_state.chats!=""):
            if(chat !=None):
                st.markdown(chat)
                st.session_state.chats=chat
                st.rerun()
            else:
                st.markdown(st.session_state.chats)
            
            
            st.session_state.files=[]
            ut.output(model,tok,st.session_state.chats)
            if(len(st.session_state.files)!=1):
                with st.sidebar:
                    st.title("Your Files")
                    for i in range(len(st.session_state.files)):
                        if(st.session_state.files[i][1]=='pdf'):
                            st.session_state.de.append(st.download_button(
                                f"{st.session_state.files[i][0]} download button",
                                data=create_pdf_bytes(st.session_state.files[i][2],st.session_state.files[i][0]),
                                file_name=f"{st.session_state.files[i][0]}.pdf",
                                key=i
                            ))
                        else:
                            st.session_state.de.append(st.download_button(
                                f"{st.session_state.files[i][0]} download button",
                                data=doc_creation(st.session_state.files[i][2],st.session_state.files[i][0]),
                                file_name=f"{st.session_state.files[i][0]}.docx",
                                key=i
                            ))
            st.session_state.chats=""
            st.rerun()
    else:
        st.rerun()
except Exception as e:
    print(e)
    st.rerun()