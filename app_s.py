import streamlit as st
import inference as inf
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import os
from transformers import AutoModelForCausalLM,AutoTokenizer,BitsAndBytesConfig,TextIteratorStreamer
import io
import tempfile
import subprocess
import threading
def chat1(model,tok,question=None):
    if(question==None):
        question=input("Give your question :: ")
    message = [
    {
        "role": "system",
        "content": """
You are a helpful assistant. You perform **only one** of the following tasks based on the user's input:

---

### 🔑 IMPORTANT TERMS

- **MULTIQUERY**: A request that includes **more than one clearly distinct task**.
    - Example: “Make a PDF on AI and another one on ML” → This is a MULTIQUERY.
    - Example: “Generate 3 PDFs on AI, ML, and DL” → This is a MULTIQUERY.
    - NOT a MULTIQUERY: “Generate a PDF on kites.” ← This is a single task.
- **Single Task**: Any prompt that contains only one action, even if the topic has subparts (e.g., “types of kites”).
- **Never guess or add tasks** not directly requested by the user.
    - ❌ Never turn “PDF on kite” into two tasks like “PDF on kite” and “PDF on types of kite.”
    - ❌ Never say “future of AI” unless user says so.

---

**STEP 1: MULTIQUERY (If the user asks for more than one task)**

If the user is asking for two or more distinct tasks (e.g., multiple documents or actions):

- STOP. Do NOT answer the question or generate content.
- Return a Python-style list of strings inside a <multiquery></multiquery> tag.

Format:

```python
<multiquery>
[
  "Task 1 in plain English.",
  "Task 2 in plain English."
]
</multiquery>

Only use this if there are **two or more explicit actions**. If only one action is present, continue to Step 2.

---

**STEP 2: DOCUMENT GENERATION**

If the user is asking for **one document only** (like "Generate a PDF on X"):
- Return a structured document with these parts:
  - A title in this format: `<title>Your Title</title>`
  - Rich, multi-paragraph content (Introduction, Details, Conclusion)
  - Exactly one of the following tags:
    - `<DownloadPDF>` for PDFs
    - `<DownloadDOCX>` for Word files
- Do NOT include <multiquery> or chatbot-like replies.

If no document is asked for, continue to Step 3.

---

**STEP 3: REGULAR RESPONSE**

If the user is asking a **general question** (e.g., "What is AI?", "Explain machine learning", "Define deep learning"):
- Give a normal, helpful answer.
- DO NOT include <multiquery>, <title>, or download tags.
- Just respond clearly and concisely.

---

**RULES:**
- Always check for MULTIQUERY first — only if there are **2+ clear tasks**.
- Never mix steps. Only one step should be followed.
- Never return content when <multiquery> is used.
- NEVER guess what the user "might also want" — only act on clear instructions.
- Never add additional tasks by yourself
"""
    },
    {
        "role": "user",
        "content": question
    }
]





    prompt=tok.apply_chat_template(message,tokenize=False)
    # print(prompt)
    tok_prompt=tok(prompt,return_tensors='pt').to('cuda')
    
    streamer = TextIteratorStreamer(tok, skip_prompt=True, skip_special_tokens=True)
    generation_kwargs = dict(tok_prompt, streamer=streamer, max_new_tokens=4096)
    thread = threading.Thread(target=model.generate, kwargs=generation_kwargs)
    thread.start()
    output_container = st.empty()
    full_output = ""
    for token in streamer:
        # print(token, end="", flush=True)
        full_output += token
        output_container.markdown(full_output)
    return full_output
    # return "Streaming done"
@st.cache_data
def doc_creation(out, title):
    docx_file = f"{title}.docx"
    doc = Document()
    doc.add_heading(title, 0)

    for line in out.strip().split("\n")[0:]:  # Skip title and blank
        if line.startswith("*"):
            para = doc.add_paragraph(line.replace("*", u"\u2022"), style='List Bullet')
        elif line.strip() == "<Download>":
            doc.add_paragraph("<Download>").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            doc.add_paragraph(line.strip())
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer.getvalue()

    
def create_pdf_bytes(out, title):
    if out is None:
        return None
    
    # Create temporary directory for file operations
    with tempfile.TemporaryDirectory() as temp_dir:
        docx_file = os.path.join(temp_dir, f"{title}.docx")
        pdf_file = os.path.join(temp_dir, f"{title}.pdf")
        
        # Create Word document (same as your original logic)
        doc = Document()
        doc.add_heading(title, 0)

        for line in out.strip().split("\n")[0:]:
            if line.startswith("*"):
                para = doc.add_paragraph(line.replace("*", u"\u2022"), style='List Bullet')
            elif line.strip() == "<Download>":
                doc.add_paragraph("<Download>").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                doc.add_paragraph(line.strip())

        # Save Word document
        doc.save(docx_file)
        subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", temp_dir,
            docx_file
        ], check=True, timeout=60)

        # Read PDF file and convert to bytes
        with open(pdf_file, 'rb') as pdf_f:
            pdf_bytes = pdf_f.read()
        
        return pdf_bytes
def clean_output(text):
    flag=-1
    ind1=text.find("<title>")
    ind2=text.find("</title>")
    title=None
    text_out=None
    # print(text[ind1+len("<think>"):ind2])
    if re.search(r"<DownloadPDF>(?![>\w])", text):
        flag=1
    elif re.search(r"<DownloadDOCX>(?![>\w])", text):
        flag=0
    elif re.search(r"</multiquery>(?![>\w])", text):
        flag=2
    if(flag!=-1):
        think=None
        title=title=text[ind1+len("<title>"):ind2]
        text_out=text[ind2+len("</title>"):]
    else:
        title=None
        text_out=text
    think=None
    return flag,think,title,text_out

def give_model():
    bnb_config = BitsAndBytesConfig(
    load_in_4bit=True,
    bnb_4bit_use_double_quant=True,
    bnb_4bit_quant_type="nf4",
    bnb_4bit_compute_dtype="float16"
    )
    tok=AutoTokenizer.from_pretrained(
        "meta-llama/Llama-3.1-8B-Instruct"
    )
    model=AutoModelForCausalLM.from_pretrained(
        "meta-llama/Llama-3.1-8B-Instruct",
        device_map='auto',
        quantization_config=bnb_config,
        # torch_dtype=torch.float32
    )
    return model,tok

st.title("Chatbot-Interface")
if "Response" not in st.session_state:
    st.session_state.Response=[]
    st.session_state.model_response=[]
# print(st.session_state)
for response in range(len(st.session_state.Response)):
    st.markdown(st.session_state.Response[response])
    st.markdown(st.session_state.model_response[response])
# chat=st.chat_input("Message...")
if 'files' not in st.session_state:
    st.session_state.files=[]
if "de" not in st.session_state:
    st.session_state.de=[]
if(len(st.session_state.de)!=0):
    for i in st.session_state.de:
        try:
            st.session_state.de.pop()
            # st.session_state.files.pop()
        except:
            pass

if(len(st.session_state.files)!=0):
    with st.sidebar:
        st.title("Your Files")
        for i in range(len(st.session_state.files)):
            if(st.session_state.files[i][1]=='pdf'):
                st.session_state.de.append(st.download_button(
                    f"{st.session_state.files[i][0]} download button",
                    data=create_pdf_bytes(st.session_state.files[i][2],st.session_state.files[i][0]),
                    file_name=f"{st.session_state.files[i][0]}.pdf",
                    key=i
                ))
            else:
                st.session_state.de.append(st.download_button(
                    f"{st.session_state.files[i][0]} download button",
                    data=doc_creation(st.session_state.files[i][2],st.session_state.files[i][0]),
                    file_name=f"{st.session_state.files[i][0]}.docx",
                    key=i
                ))

if 'model' not in st.session_state:
    st.session_state.model = None
if 'tok' not in st.session_state:
    st.session_state.tok = None

# Load model only once
if st.session_state.model is None:
    with st.spinner("Loading model... This may take a few minutes."):
        st.session_state.model, st.session_state.tok = give_model()
    st.success("Model loaded successfully!")

# Use the model
model = st.session_state.model
tok = st.session_state.tok
chat=st.chat_input("Message...")
# def trunc_history()

# def generate_history():
#     history=[]
#     for i in range(len(st.session_state.Response)):
#         di1={"role":"user","content":st.session_state.Response[i]}
#         di2={"role":'assistant',"content":st.session_state.model_response[i]}
#         history.append(di1)
#         history.append(di2)
#     return history


def return_list(text):
    ind1=text.find("<multiquery>")
    ind2=text.find("</multiquery")
    return eval(text[ind1+len("<multiquery>"):ind2].strip())

def output(model,tok,question=None):
    # hist1=generate_history()
    out=chat1(model,tok,question)
    flag,think,title,text=clean_output(out)
    try:
        if(flag==-1):
            print("Flag",flag)
            st.markdown(text)
        elif(flag==2):
            li=return_list(text)
            for i in li:
                output(model,tok,i)
                # st.session_state.files.pop()
        else:
            st.markdown("Word" if flag==0 else "PDF")
            st.markdown("-"*50)
            st.markdown(think)
            st.markdown("-"*50)
            st.markdown(title)
            st.markdown("-"*50)
            st.markdown(text)
            st.markdown('-'*50)
            if(flag==0):
                st.session_state.files.append([title,'docx',text])
        #     if(flag==0):
        #         st.session_state.files.append(st.download_button(
        #             label=f"{title} download button",
        #             data=doc_creation(text,title),
        #             file_name=f"{title}.docx",
        #             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        #         ))
            elif(flag==1):
                st.session_state.files.append([title,'pdf',text])
        #         st.session_state.files.append(st.download_button(
        #             label=f"{title} download button",
        #             data=create_pdf_bytes(text,title) ,
        #             file_name=f"{title}.pdf",
        #             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        #         ))
        st.session_state.Response.append(question)
        st.session_state.model_response.append(text)
    except Exception as e:
        print("Error")
        print(e)
        st.markdown("Word" if flag==0 else "PDF")
        st.markdown("-"*50)
        st.markdown(think)
        st.markdown("-"*50)
        st.markdown(title)
        st.markdown("-"*50)
        st.markdown(text)
        st.markdown('-'*50)
if 'chats' not in st.session_state:
    st.session_state.chats=""
# print("*"*200)
# print(chat)
# print(st.session_state.chats)
# print("*"*200)
if(chat!=None or st.session_state.chats!=""):
    if(chat !=None):
        st.markdown(chat)
        st.session_state.chats=chat
        st.rerun()
    else:
        st.markdown(st.session_state.chats)
    
    
    st.session_state.files=[]
    # hist=generate_history()
    output(model,tok,st.session_state.chats)
    if(len(st.session_state.files)!=1):
        with st.sidebar:
            st.title("Your Files")
            for i in range(len(st.session_state.files)):
                if(st.session_state.files[i][1]=='pdf'):
                    st.session_state.de.append(st.download_button(
                        f"{st.session_state.files[i][0]} download button",
                        data=create_pdf_bytes(st.session_state.files[i][2],st.session_state.files[i][0]),
                        file_name=f"{st.session_state.files[i][0]}.pdf",
                        key=i
                    ))
                else:
                    st.session_state.de.append(st.download_button(
                        f"{st.session_state.files[i][0]} download button",
                        data=doc_creation(st.session_state.files[i][2],st.session_state.files[i][0]),
                        file_name=f"{st.session_state.files[i][0]}.docx",
                        key=i
                    ))
    st.session_state.chats=""
    st.rerun()
