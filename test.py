import re
import io
import tempfile
import subprocess
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import shutil
text="""
assistant

**<title>How AirPods Work</title>**

AirPods are a popular line of wireless earbuds designed and marketed by Apple Inc. They use a combination of Bluetooth technology, microprocessors, and software to provide a seamless listening experience. Here's a detailed explanation of how they work:

### Components and Architecture

AirPods consist of several key components:

1. **Microphones**: AirPods have two microphones, one on each earbud, which pick up sound from the environment and send it to the connected device.
2. **Speakers**: Each earbud has a speaker that converts digital audio signals into sound waves.
3. **Bluetooth Chip**: The Bluetooth chip in AirPods enables wireless connectivity to devices such as iPhones, iPads, and Macs.
4. **Battery**: AirPods have a rechargeable battery that powers the earbuds and provides up to 5 hours of listening time on a single charge.
5. **Controller**: The earbuds have a small controller that allows users to control playback, volume, and other functions.

### Connectivity and Pairing

When you first set up AirPods, they pair with your device via Bluetooth. The pairing process involves:

1. **Discovery**: The AirPods discover nearby devices and display a list of available options.
2. **Authentication**: The user selects their device from the list, and the AirPods authenticate with the device.
3. **Connection**: The AirPods establish a secure connection with the device, enabling wireless audio streaming.

### Audio Signal Processing

AirPods use a combination of hardware and software to process audio signals:

1. **Audio Signal Reception**: The microphones receive sound waves from the environment and send them to the connected device.
2. **Noise Cancellation**: AirPods use noise cancellation technology to reduce background noise and improve sound quality.
3. **Audio Signal Processing**: The device processes the audio signal and sends it to the AirPods for playback.

### Additional Features

AirPods also offer several additional features, including:

1. **Siri Integration**: AirPods allow users to access Siri, Apple's virtual assistant, with voice commands.
2. **Wireless Charging**: AirPods support wireless charging, making it easy to keep them charged without the hassle of cables.
3. **Seamless Handoff**: AirPods enable seamless handoff between devices, allowing users to switch between listening on different devices.

**<DownloadPDF>**

---

**<title>The Benefits of AirPods</title>**

AirPods have revolutionized the way we listen to music and podcasts on-the-go. Here are some of the key benefits of using AirPods:

### Convenience

1. **Wireless Freedom**: AirPods offer wireless freedom, allowing users to move around without being tethered to their device.
2. **Easy Pairing**: AirPods pair quickly and easily with devices, making it simple to get started.
3. **Seamless Handoff**: AirPods enable seamless handoff between devices, allowing users to switch between listening on different devices.

### Audio Quality

1. **Crystal-Clear Sound**: AirPods provide high-quality audio, with clear and crisp sound that's perfect for listening to music and podcasts.
2. **Noise Cancellation**: AirPods use noise cancellation technology to reduce background noise and improve sound quality.
3. **Customizable EQ**: AirPods allow users to customize the EQ settings to suit their listening preferences.

### Health and Safety

1. **Ear Health**: AirPods are designed to promote ear health, with features like volume limits and earbud tips that help prevent ear damage.
2. **Safety**: AirPods are designed with safety in mind, with features like automatic shut-off and water resistance.

### Additional Benefits

1. **Long Battery Life**: AirPods have a long battery life, with up to 5 hours of listening time on a single charge.
2. **Wireless Charging**: AirPods support wireless charging, making it easy to keep them charged without the hassle of cables.
3. **Integration with Apple Devices**: AirPods integrate seamlessly with Apple devices, allowing users to access features like Siri and iCloud.

**<DownloadDOCX>**


**<title>How AirPods Work</title>**

AirPods are a popular line of wireless earbuds designed and marketed by Apple Inc. They use a combination of Bluetooth technology, microprocessors, and software to provide a seamless listening experience. Here's a detailed explanation of how they work:

### Components and Architecture

AirPods consist of several key components:

1. **Microphones**: AirPods have two microphones, one on each earbud, which pick up sound from the environment and send it to the connected device.
2. **Speakers**: Each earbud has a speaker that converts digital audio signals into sound waves.
3. **Bluetooth Chip**: The Bluetooth chip in AirPods enables wireless connectivity to devices such as iPhones, iPads, and Macs.
4. **Battery**: AirPods have a rechargeable battery that powers the earbuds and provides up to 5 hours of listening time on a single charge.
5. **Controller**: The earbuds have a small controller that allows users to control playback, volume, and other functions.

### Connectivity and Pairing

When you first set up AirPods, they pair with your device via Bluetooth. The pairing process involves:

1. **Discovery**: The AirPods discover nearby devices and display a list of available options.
2. **Authentication**: The user selects their device from the list, and the AirPods authenticate with the device.
3. **Connection**: The AirPods establish a secure connection with the device, enabling wireless audio streaming.

### Audio Signal Processing

AirPods use a combination of hardware and software to process audio signals:

1. **Audio Signal Reception**: The microphones receive sound waves from the environment and send them to the connected device.
2. **Noise Cancellation**: AirPods use noise cancellation technology to reduce background noise and improve sound quality.
3. **Audio Signal Processing**: The device processes the audio signal and sends it to the AirPods for playback.

### Additional Features

AirPods also offer several additional features, including:

1. **Siri Integration**: AirPods allow users to access Siri, Apple's virtual assistant, with voice commands.
2. **Wireless Charging**: AirPods support wireless charging, making it easy to keep them charged without the hassle of cables.
3. **Seamless Handoff**: AirPods enable seamless handoff between devices, allowing users to switch between listening on different devices.

**<DownloadPDF>**

---

**<title>The Benefits of AirPods</title>**

AirPods have revolutionized the way we listen to music and podcasts on-the-go. Here are some of the key benefits of using AirPods:

### Convenience

1. **Wireless Freedom**: AirPods offer wireless freedom, allowing users to move around without being tethered to their device.
2. **Easy Pairing**: AirPods pair quickly and easily with devices, making it simple to get started.
3. **Seamless Handoff**: AirPods enable seamless handoff between devices, allowing users to switch between listening on different devices.

### Audio Quality

1. **Crystal-Clear Sound**: AirPods provide high-quality audio, with clear and crisp sound that's perfect for listening to music and podcasts.
2. **Noise Cancellation**: AirPods use noise cancellation technology to reduce background noise and improve sound quality.
3. **Customizable EQ**: AirPods allow users to customize the EQ settings to suit their listening preferences.

### Health and Safety

1. **Ear Health**: AirPods are designed to promote ear health, with features like volume limits and earbud tips that help prevent ear damage.
2. **Safety**: AirPods are designed with safety in mind, with features like automatic shut-off and water resistance.

### Additional Benefits

1. **Long Battery Life**: AirPods have a long battery life, with up to 5 hours of listening time on a single charge.
2. **Wireless Charging**: AirPods support wireless charging, making it easy to keep them charged without the hassle of cables.
3. **Integration with Apple Devices**: AirPods integrate seamlessly with Apple devices, allowing users to access features like Siri and iCloud.

"""
def clean_output(text):
    parts = re.split(r'(<title> | <DownloadPDF> | <DownloadDOCX>)', text)
    out=[]
    text=""
    for i in range(len(parts)):
        print(parts[i])
        print()
        print('-'*200)
        print()
        # if ('<title>' in parts[i] and (parts[i]!="<DownloadPDF>" and parts[i]!="<DownloadDOCX>")):
        #     try:
        #         ind=parts[i].find('<title>')
        #         ind2=parts[i][ind+len('<title>'):].find('</title>')
        #         if(parts[i+1]=='<DownloadDOCX>'):
        #             li=[1,parts[i][ind+len("<title>"):ind+len("<title>")+ind2],parts[i]]
        #             out.append(li)
        #         elif(parts[i+1]=='<DownloadPDF>'):
        #             li=[2,parts[i][ind+len("<title>"):ind+len("<title>")+ind2],parts[i]]
        #             out.append(li)
        #         else:
        #             li=[0,parts[i]]
        #             out.append(li)
        #     except:
        #         li=[0,parts[i]]
        #         out.append(li)
        #         pass
        # else:
        #     if(parts[i]!="<DownloadPDF>" and parts[i]!="<DownloadDOCX>"):
        #         li=[0,parts[i]]
        #         out.append(li)
        #     print(parts[i])
    return out,text
t,t1= clean_output(text)

for i in t:
    print(i)
    print()
print()
print(t1)
# def add_bold_paragraph(doc, text, style=None):
#     """Add a paragraph where **text** is made bold."""
#     para = doc.add_paragraph(style=style)
#     parts = re.split(r"(\*\*.*?\*\*)", text)
#     for part in parts:
#         if part.startswith("**") and part.endswith("**"):
#             run = para.add_run(part[2:-2])
#             run.bold = True
#         else:
#             para.add_run(part)
#     return para

# from docx import Document
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# import os, tempfile, subprocess, shutil
# import re

# def add_bold_paragraph(doc, text, style=None):
#     para = doc.add_paragraph(style=style) if style else doc.add_paragraph()
#     parts = re.split(r"(\*\*.*?\*\*)", text)
#     for part in parts:
#         if part.startswith("**") and part.endswith("**"):
#             run = para.add_run(part[2:-2])
#             run.bold = True
#         else:
#             para.add_run(part)
#     return para

# def doc_creation(raw_text):
#     title_match = re.search(r"<title>(.*?)</title>", raw_text, re.DOTALL)
#     title = title_match.group(1).strip() if title_match else "Document"

#     with tempfile.TemporaryDirectory() as tmp:
#         docx_path = os.path.join(tmp, f"{title}.docx")
#         pdf_path  = os.path.join(tmp, f"{title}.pdf")
#         doc = Document()
#         doc.add_heading(title, level=0)
#         body = re.sub(r"^.*?</title>\s*", "", raw_text, flags=re.DOTALL)
#         lines = body.splitlines()

#         for raw in lines:
#             line = raw.strip()
#             if not line or line in ("---", "**"):
#                 continue
#             if line.startswith("### "):
#                 doc.add_heading(line[4:].strip(), level=1)
#             elif re.match(r"^\d+\.\s+.*", line):
#                 add_bold_paragraph(doc, line)
#             elif m := re.match(r"^[*•]\s*(.*)", line):
#                 content = m.group(1)
#                 add_bold_paragraph(doc, content, style='List Bullet')
#             elif line in ("<DownloadPDF>", "<DownloadDOCX>"):
#                 p = doc.add_paragraph(line)
#                 p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#             else:
#                 add_bold_paragraph(doc, line)
#     doc_buffer = io.BytesIO()
#     doc.save(doc_buffer)
#     doc_buffer.seek(0)
#     return doc_buffer.getvalue()
# def markdown_to_pdf(raw_text: str):
#     title_match = re.search(r"<title>(.*?)</title>", raw_text, re.DOTALL)
#     title = title_match.group(1).strip() if title_match else "Document"

#     with tempfile.TemporaryDirectory() as tmp:
#         docx_path = os.path.join(tmp, f"{title}.docx")
#         pdf_path  = os.path.join(tmp, f"{title}.pdf")
#         doc = Document()
#         doc.add_heading(title, level=0)
#         body = re.sub(r"^.*?</title>\s*", "", raw_text, flags=re.DOTALL)
#         lines = body.splitlines()

#         for raw in lines:
#             line = raw.strip()
#             if not line or line in ("---", "**"):
#                 continue
#             if line.startswith("### "):
#                 doc.add_heading(line[4:].strip(), level=1)
#             elif re.match(r"^\d+\.\s+.*", line):
#                 add_bold_paragraph(doc, line)
#             elif m := re.match(r"^[*•]\s*(.*)", line):
#                 content = m.group(1)
#                 add_bold_paragraph(doc, content, style='List Bullet')
#             elif line in ("<DownloadPDF>", "<DownloadDOCX>"):
#                 p = doc.add_paragraph(line)
#                 p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#             else:
#                 add_bold_paragraph(doc, line)
#         doc.save(docx_path)
#         subprocess.run([
#             "libreoffice", "--headless",
#             "--convert-to", "pdf", "--outdir", tmp,
#             docx_path
#         ], check=True, timeout=60)

#         with open(pdf_path, 'rb') as pdf_f:
#             pdf_bytes = pdf_f.read()
        
#         return pdf_bytes

# print(doc_creation(t[1][2],t[1][1]))